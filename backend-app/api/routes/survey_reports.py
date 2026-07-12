import os
import io
import uuid
import base64
import datetime
import re
import copy
import docx
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional, List
from loguru import logger
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage

from core.database import get_db
from core.config import settings
from models.domain import Project, SurveyReport, SurveyReportSubcontractor, SurveyReportWorkScope
from services.ai.agents import survey_report_agent
from services.azure_storage import azure_storage

router = APIRouter()

UPLOAD_DIR = os.path.abspath(os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads"))
TEMPLATES_DIR = os.path.abspath(os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "SurveyTemplate"))

if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

class SubcontractorDTO(BaseModel):
    scope_of_works: str
    subcontractor: str

class WorkScopeDTO(BaseModel):
    title: str
    description: Optional[str] = None
    image_base64: Optional[str] = None
    status: Optional[str] = "New"
    start_date: Optional[str] = ""
    categories: Optional[str] = ""  # Comma-separated categories
    permits: Optional[str] = ""     # Comma-separated permits

class SurveyReportGeneratePayload(BaseModel):
    project_id: str
    template_id: str
    document_name: Optional[str] = None
    doc_nr: Optional[str] = None
    reference: Optional[str] = None
    revision: Optional[str] = None
    company_name: Optional[str] = None
    vessel_name: Optional[str] = None
    arrival_date: Optional[str] = None  # YYYY-MM-DD
    total_lead_time: Optional[int] = None
    drydock_duration: Optional[int] = None
    cover_image_base64: Optional[str] = None
    interior_image_base64: Optional[str] = None
    company_logo_base64: Optional[str] = None
    subcontractors: List[SubcontractorDTO]
    shipyard_scopes: List[WorkScopeDTO]

# Helper to save base64 string to a local file or Azure Blob Storage
def save_base64_image(base64_str, folder):
    if not base64_str:
        return None
    try:
        if "," in base64_str:
            base64_str = base64_str.split(",")[1]
        img_data = base64.b64decode(base64_str)
        filename = f"img_{uuid.uuid4()}.jpg"
        
        use_azure = bool(settings.AZURE_STORAGE_CONNECTION_STRING or (settings.AZURE_STORAGE_ACCOUNT_NAME and settings.AZURE_STORAGE_ACCOUNT_KEY))
        if use_azure:
            try:
                logger.info(f"Uploading base64 image '{filename}' to Azure Storage container: {settings.AZURE_STORAGE_CONTAINER_IMAGES}")
                url = azure_storage.upload_blob(
                    container_name=settings.AZURE_STORAGE_CONTAINER_IMAGES,
                    blob_name=filename,
                    data=img_data
                )
                return url
            except Exception as e:
                logger.error(f"Error uploading image to Azure Storage: {str(e)}. Falling back to local storage.")
        
        path = os.path.join(folder, filename)
        with open(path, "wb") as f:
            f.write(img_data)
        return path
    except Exception as e:
        logger.error(f"Error saving base64 image: {str(e)}")
        return None

# Helper to download image bytes from Azure Blob URL or local file path
def get_image_bytes(path_or_url: str) -> Optional[bytes]:
    if not path_or_url:
        return None
    if path_or_url.startswith("http://") or path_or_url.startswith("https://"):
        try:
            import urllib.parse
            # Extract blob name from URL (the last segment)
            blob_name = path_or_url.split("/")[-1]
            if "?" in blob_name:
                blob_name = blob_name.split("?")[0]
            blob_name = urllib.parse.unquote(blob_name)
            
            logger.info(f"Downloading image '{blob_name}' from Azure Storage container: {settings.AZURE_STORAGE_CONTAINER_IMAGES}")
            return azure_storage.download_blob(settings.AZURE_STORAGE_CONTAINER_IMAGES, blob_name)
        except Exception as e:
            logger.error(f"Failed to download image from Azure Blob {path_or_url}: {e}")
            # Fallback to standard HTTP download in case it's an external URL
            try:
                import requests
                resp = requests.get(path_or_url, timeout=5)
                if resp.status_code == 200:
                    return resp.content
            except Exception as re:
                logger.error(f"Failed to fetch image via HTTP {path_or_url}: {re}")
            return None
    else:
        # It is a local file path
        if os.path.exists(path_or_url):
            try:
                with open(path_or_url, "rb") as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Failed to read local image {path_or_url}: {e}")
        return None

@router.get("/templates")
def get_templates():
    """
    Returns lists of available survey report templates and their schemas.
    """
    templates_list = []
    use_azure = bool(settings.AZURE_STORAGE_CONNECTION_STRING or (settings.AZURE_STORAGE_ACCOUNT_NAME and settings.AZURE_STORAGE_ACCOUNT_KEY))
    files = []
    
    if use_azure:
        try:
            logger.info(f"Listing templates from Azure Storage container: {settings.AZURE_STORAGE_CONTAINER_TEMPLATES}")
            all_blobs = azure_storage.list_blobs(settings.AZURE_STORAGE_CONTAINER_TEMPLATES)
            files = [f for f in all_blobs if f.endswith(".docx")]
        except Exception as e:
            logger.error(f"Error fetching templates from Azure Storage: {str(e)}. Falling back to local templates.")
            use_azure = False
            
    if not use_azure:
        if os.path.exists(TEMPLATES_DIR):
            files = [f for f in os.listdir(TEMPLATES_DIR) if f.endswith(".docx")]
            
    for f in files:
        if f == "Intermediate Survey Template.docx":
            name = "Intermediate Survey Template"
        elif f == "0318 -Intermediate Survey - 2026 - Rev_draft (1).docx":
            name = "0318 - Intermediate Survey - 2026 Draft"
        else:
            name = f.replace(".docx", "")
        templates_list.append({
            "id": f,
            "name": name,
            "filename": f,
            "fields": [
                {"name": "company_name", "type": "string", "label": "Company Name", "default": "Van Oord Marine Ingenuity"},
                {"name": "vessel_name", "type": "string", "label": "Vessel Name", "default": "HAM 318"},
                {"name": "document_name", "type": "string", "label": "Document Name", "default": "0318 - Intermediate Survey - 2026"},
                {"name": "doc_nr", "type": "string", "label": "Document Number", "default": "VOMS-PR3.07-SMD-IN-01-05"},
                {"name": "reference", "type": "string", "label": "Reference Number", "default": "VOMS-PR3.07-SMD-IN-01"},
                {"name": "revision", "type": "string", "label": "Revision", "default": "1"},
                {"name": "arrival_date", "type": "date", "label": "Arrival Date", "default": "2026-06-01"},
                {"name": "total_lead_time", "type": "number", "label": "Total Lead Time (Days)", "default": 38},
                {"name": "drydock_duration", "type": "number", "label": "Drydock Duration (Days)", "default": 21},
                {"name": "cover_image", "type": "image", "label": "Vessel Cover Image"},
                {"name": "interior_image", "type": "image", "label": "Vessel Layout/Interior Image"}
            ]
        })
    return {"templates": templates_list}

@router.get("/templates/{filename}/download")
def download_template(filename: str):
    """
    Serves raw Word templates directly.
    """
    use_azure = bool(settings.AZURE_STORAGE_CONNECTION_STRING or (settings.AZURE_STORAGE_ACCOUNT_NAME and settings.AZURE_STORAGE_ACCOUNT_KEY))
    
    if use_azure:
        try:
            logger.info(f"Downloading template '{filename}' from Azure Storage container: {settings.AZURE_STORAGE_CONTAINER_TEMPLATES}")
            template_bytes = azure_storage.download_blob(settings.AZURE_STORAGE_CONTAINER_TEMPLATES, filename)
            return StreamingResponse(
                io.BytesIO(template_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        except Exception as e:
            logger.error(f"Error downloading template from Azure Storage: {str(e)}. Falling back to local templates.")
            use_azure = False
            
    if not use_azure:
        file_path = os.path.join(TEMPLATES_DIR, filename)
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Template file not found")
        return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)

@router.post("/generate-ai-description")
async def generate_ai_description(
    file: UploadFile = File(...),
    prompt_caption: str = Form(""),
):
    """
    Uses SurveyReportAgent to generate a short, crispy shipyard work description from a photo.
    """
    try:
        content = await file.read()
        encoded_image = base64.b64encode(content).decode("utf-8")

        state = {
            "image_base64": encoded_image,
            "prompt_caption": prompt_caption
        }
        result = survey_report_agent.invoke(state)
        desc_text = result.get("description", "")
        return {"description": desc_text}
    except Exception as e:
        logger.error(f"Error generating description with SurveyReportAgent: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Vision AI generation failed: {str(e)}")

def replace_text_in_paragraph(p, key, val):
    if key not in p.text:
        return False
    replaced = False
    for run in p.runs:
        if key in run.text:
            t_nodes = run._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            for t_node in t_nodes:
                if t_node.text and key in t_node.text:
                    t_node.text = t_node.text.replace(key, val)
            replaced = True
    if not replaced:
        t_nodes = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        if t_nodes:
            full_text = "".join(t.text for t in t_nodes if t.text)
            if key in full_text:
                replaced_text = full_text.replace(key, val)
                t_nodes[0].text = replaced_text
                for t in t_nodes[1:]:
                    t.text = ""
                replaced = True
    return replaced

def replace_in_element(element, doc, replacements):
    from docx.text.paragraph import Paragraph
    p_elements = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    for p_el in p_elements:
        p = Paragraph(p_el, doc)
        for key, val in replacements.items():
            replace_text_in_paragraph(p, key, val)

def replace_placeholder_with_image(p, key, img_path_or_bytes, width=None):
    if key not in p.text:
        return False
        
    def get_image_file():
        if isinstance(img_path_or_bytes, bytes):
            return io.BytesIO(img_path_or_bytes)
        return img_path_or_bytes

    for run in p.runs:
        if key in run.text:
            t_nodes = run._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            for t_node in t_nodes:
                if t_node.text and key in t_node.text:
                    t_node.text = t_node.text.replace(key, "")
            run.add_picture(get_image_file(), width=width)
            return True
    t_nodes = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    if t_nodes:
        full_text = "".join(t.text for t in t_nodes if t.text)
        if key in full_text:
            first_t = t_nodes[0]
            replaced_text = full_text.replace(key, "")
            first_t.text = replaced_text
            for t in t_nodes[1:]:
                t.text = ""
            
            parent_r = None
            curr = first_t
            for _ in range(5):
                if curr is None:
                    break
                parent = curr.getparent()
                if parent is not None and parent.tag.endswith('r'):
                    parent_r = parent
                    break
                curr = parent
            
            if parent_r is not None:
                from docx.text.run import Run
                run = Run(parent_r, p)
                run.add_picture(get_image_file(), width=width)
                return True
    return False

def replace_logo_in_element(element, doc, logo_path_or_bytes, width):
    from docx.text.paragraph import Paragraph
    p_elements = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    for p_el in p_elements:
        p = Paragraph(p_el, doc)
        if "companyNameLogo" in p.text:
            replace_placeholder_with_image(p, "companyNameLogo", logo_path_or_bytes, width)

@router.post("/generate")
def generate_report(payload: SurveyReportGeneratePayload, db: Session = Depends(get_db)):
    """
    Stores survey report details in DB and compiles the Word (DOCX) document from the template.
    """
    try:
        try:
            project_uuid = uuid.UUID(payload.project_id)
        except ValueError:
            logger.warning(f"Invalid UUID received for project_id: {payload.project_id}. Attempting fallback...")
            project = db.query(Project).first()
            if not project:
                project = Project(
                    project_code="PRJ-DEMO",
                    project_title="Demo Project (Auto-created)",
                    project_type="Demo",
                    vessel_name="HAM 318",
                    vessel_type="Dredger"
                )
                db.add(project)
                db.flush()
            project_uuid = project.project_id

        project = db.query(Project).filter(Project.project_id == project_uuid).first()
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")

        # 1. Save base64 images
        cover_path = save_base64_image(payload.cover_image_base64, UPLOAD_DIR)
        interior_path = save_base64_image(payload.interior_image_base64, UPLOAD_DIR)
        logo_path = save_base64_image(payload.company_logo_base64, UPLOAD_DIR)

        # 2. Insert SurveyReport record
        arrival_date_val = None
        if payload.arrival_date:
            arrival_date_val = datetime.datetime.strptime(payload.arrival_date, "%Y-%m-%d").date()

        report = SurveyReport(
            project_id=project_uuid,
            template_id=payload.template_id,
            document_name=payload.document_name or "Intermediate Survey Report",
            doc_nr=payload.doc_nr,
            reference=payload.reference,
            revision=payload.revision,
            company_name=payload.company_name or "Van Oord Marine Ingenuity",
            vessel_name=payload.vessel_name or "HAM 318",
            arrival_date=arrival_date_val,
            total_lead_time=payload.total_lead_time or 38,
            drydock_duration=payload.drydock_duration or 21,
            cover_image_path=cover_path,
            interior_image_path=interior_path,
            company_logo_path=logo_path,
            status="Draft"
        )
        db.add(report)
        db.flush()

        # 3. Add Subcontractors
        for sub in payload.subcontractors:
            db_sub = SurveyReportSubcontractor(
                report_id=report.report_id,
                scope_of_works=sub.scope_of_works,
                subcontractor=sub.subcontractor
            )
            db.add(db_sub)

        # 4. Add Shipyard Scopes
        scopes_list = []
        for idx, scope in enumerate(payload.shipyard_scopes):
            scope_img_path = save_base64_image(scope.image_base64, UPLOAD_DIR)
            db_scope = SurveyReportWorkScope(
                report_id=report.report_id,
                sequence_number=idx + 1,
                title=scope.title,
                description=scope.description,
                image_path=scope_img_path,
                status=scope.status or "New",
                start_date=scope.start_date or "",
                categories=scope.categories or "",
                permits=scope.permits or ""
            )
            db.add(db_scope)
            scopes_list.append(db_scope)

        db.commit()
        db.refresh(report)

        # 5. Populate and generate DOCX file from template
        use_azure = bool(settings.AZURE_STORAGE_CONNECTION_STRING or (settings.AZURE_STORAGE_ACCOUNT_NAME and settings.AZURE_STORAGE_ACCOUNT_KEY))
        template_bytes = None
        
        if use_azure:
            try:
                logger.info(f"Loading template '{payload.template_id}' from Azure Storage container: {settings.AZURE_STORAGE_CONTAINER_TEMPLATES}")
                template_bytes = azure_storage.download_blob(settings.AZURE_STORAGE_CONTAINER_TEMPLATES, payload.template_id)
            except Exception as e:
                logger.error(f"Error loading template from Azure Storage: {str(e)}. Falling back to local filesystem.")
                use_azure = False
                
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        if template_bytes is not None:
            doc = Document(io.BytesIO(template_bytes))
        else:
            template_file_path = os.path.join(TEMPLATES_DIR, payload.template_id)
            if not os.path.exists(template_file_path):
                raise HTTPException(status_code=400, detail="Selected template file does not exist")
            doc = Document(template_file_path)

        # Define replacements dictionary
        current_date_str = datetime.date.today().strftime("%d %B %Y")
        replacements = {
            "companyNameLogo": report.company_name,
            "vesselName": report.vessel_name,
            "documentNumber": report.doc_nr or "",
            "referenceNumber": report.reference or "",
            "documentType": report.document_name or "",
            "arrivalDate": report.arrival_date.strftime("%d-%m-%Y") if report.arrival_date else "",
            "totalLeadTimeDays": str(report.total_lead_time),
            "dryDockDurationDays": str(report.drydock_duration),
            
            # Map hardcoded draft values from the 0318 -Intermediate Survey - 2026 - Rev_draft (1).docx template
            "0326 – Maintenance & Repairs 2025": report.document_name or "",
            "0326 - Maintenance & Repairs 2025": report.document_name or "",
            "0326 \u2013 Maintenance and Repairs 2025": report.document_name or "",
            "VOMS-PR3.07-SMD-IN-01-05": report.doc_nr or "",
            "VOMS-PR3.07-SMD-IN-01": report.reference or "",
            "5 February 2026": current_date_str,
            "date": current_date_str
        }

        # Replace Company Logo if provided, otherwise it will fall back to text replacement via replacements dict
        logo_bytes = get_image_bytes(logo_path)
        if logo_bytes:
            replace_logo_in_element(doc.element.body, doc, logo_bytes, Inches(1.8))
            for section in doc.sections:
                for h_type in ['header', 'first_page_header', 'even_page_header']:
                    header = getattr(section, h_type, None)
                    if header:
                        replace_logo_in_element(header._element, doc, logo_bytes, Inches(1.2))
                for f_type in ['footer', 'first_page_footer', 'even_page_footer']:
                    footer = getattr(section, f_type, None)
                    if footer:
                        replace_logo_in_element(footer._element, doc, logo_bytes, Inches(1.2))

        # Text Replacements in body and headers/footers recursively (to cover drawings, cells, shapes, headers, footers)
        replace_in_element(doc.element.body, doc, replacements)
        for section in doc.sections:
            for h_type in ['header', 'first_page_header', 'even_page_header']:
                header = getattr(section, h_type, None)
                if header:
                    replace_in_element(header._element, doc, replacements)
            for f_type in ['footer', 'first_page_footer', 'even_page_footer']:
                footer = getattr(section, f_type, None)
                if footer:
                    replace_in_element(footer._element, doc, replacements)

        # Helper to set cell shading (background color)
        def set_cell_background(cell, color_hex):
            shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
            cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

        # Insert Cover Image
        cover_bytes = get_image_bytes(cover_path)
        if cover_bytes:
            for p in doc.paragraphs:
                if "vesselCoverageImage" in p.text or "vesselCoverImage" in p.text:
                    replace_placeholder_with_image(p, "vesselCoverageImage", cover_bytes, width=Inches(5.0))
                    replace_placeholder_with_image(p, "vesselCoverImage", cover_bytes, width=Inches(5.0))
                    break

        # Insert Interior Layout Image
        interior_bytes = get_image_bytes(interior_path)
        if interior_bytes:
            for p in doc.paragraphs:
                if "vesselInteriorImage" in p.text:
                    replace_placeholder_with_image(p, "vesselInteriorImage", interior_bytes, width=Inches(5.0))
                    break

        # Replace subcontractors list paragraph
        for p in doc.paragraphs:
            if "subContractorsName" in p.text:
                p.text = ""
                for idx, sub in enumerate(payload.subcontractors):
                    p.text += f"{sub.subcontractor}:\t{sub.scope_of_works}\n"
                p.text = p.text.rstrip()
                break

        # Fill Subcontractors table (Table 7)
        sub_table = None
        for t in doc.tables:
            if len(t.columns) == 3 and "Scope of works" in t.cell(0, 1).text:
                sub_table = t
                break

        if sub_table:
            total_rows = len(sub_table.rows)
            for idx, sub in enumerate(payload.subcontractors):
                row_idx = idx + 1
                if row_idx < total_rows:
                    row = sub_table.rows[row_idx]
                    row.cells[0].text = str(row_idx)
                    row.cells[1].text = sub.scope_of_works
                    row.cells[2].text = sub.subcontractor
                else:
                    row = sub_table.add_row()
                    row.cells[0].text = str(row_idx)
                    row.cells[1].text = sub.scope_of_works
                    row.cells[2].text = sub.subcontractor
            # Remove empty rows
            unused_start = len(payload.subcontractors) + 1
            while len(sub_table.rows) > unused_start:
                tbl_element = sub_table._tbl
                tr_element = sub_table.rows[-1]._tr
                tbl_element.remove(tr_element)

        # Fill Index table (Table 2)
        index_table = doc.tables[2]
        total_index_rows = len(index_table.rows)
        for idx, scope in enumerate(scopes_list):
            row_idx = idx + 1
            ref = f"4.{row_idx}"
            if row_idx < total_index_rows:
                row = index_table.rows[row_idx]
                row.cells[0].text = str(row_idx)
                row.cells[1].text = scope.title
                row.cells[2].text = ref
                row.cells[3].text = scope.status or "New"
                row.cells[4].text = scope.start_date or ""
            else:
                row = index_table.add_row()
                row.cells[0].text = str(row_idx)
                row.cells[1].text = scope.title
                row.cells[2].text = ref
                row.cells[3].text = scope.status or "New"
                row.cells[4].text = scope.start_date or ""

        unused_index_start = len(scopes_list) + 1
        while len(index_table.rows) > unused_index_start:
            tbl_element = index_table._tbl
            tr_element = index_table.rows[-1]._tr
            tbl_element.remove(tr_element)

        # Cloned Shipyard Scopes block processing
        body = doc.element.body
        placeholder_idx = -1
        for i, child in enumerate(body):
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                p = docx.text.paragraph.Paragraph(child, doc)
                if 'shipyard_scopes' in p.text:
                    placeholder_idx = i
                    break

        if placeholder_idx != -1:
            tbl3_xml = None
            tbl4_xml = None
            p_draw_xml = None
            p_folder_xml = None
            tbl5_xml = None
            tbl6_xml = None
            block_end_idx = -1

            for i in range(placeholder_idx + 1, len(body)):
                child = body[i]
                tag = child.tag.split('}')[-1]
                if tag == 'p':
                    p = docx.text.paragraph.Paragraph(child, doc)
                    if 'Overview Subcontractors' in p.text:
                        block_end_idx = i
                        break
                    if 'Drawings/ documents related:' in p.text:
                        p_draw_xml = copy.deepcopy(child)
                    elif 'See folder 4.' in p.text:
                        p_folder_xml = copy.deepcopy(child)
                elif tag == 'tbl':
                    t = docx.table.Table(child, doc)
                    first_cell_text = t.cell(0, 0).text.strip()
                    if 'Steel' in first_cell_text:
                        tbl3_xml = copy.deepcopy(child)
                    elif 'Confined space' in first_cell_text:
                        tbl4_xml = copy.deepcopy(child)
                    elif 'Job description for Yard' in first_cell_text:
                        tbl5_xml = copy.deepcopy(child)
                    elif 'Scope of work related to this job done by Crew' in first_cell_text:
                        tbl6_xml = copy.deepcopy(child)

            if block_end_idx != -1:
                # Remove the original template placeholder scope items
                for idx in range(block_end_idx - 1, placeholder_idx - 1, -1):
                    body.remove(body[idx])

                # Insert populated scopes
                insert_idx = placeholder_idx
                for idx, scope in enumerate(scopes_list):
                    # Title
                    p_title = doc.add_paragraph(style='Heading 2')
                    p_title_run = p_title.add_run(f"4.{idx+1} {scope.title}")
                    p_title_run.bold = True
                    p_title_run.font.name = "Arial"
                    p_title_run.font.size = Pt(14)
                    p_title_run.font.color.rgb = RGBColor(30, 58, 138)
                    body.insert(insert_idx, p_title._p)
                    insert_idx += 1

                    # Checklist categories (Table 3)
                    t3 = docx.table.Table(copy.deepcopy(tbl3_xml), doc)
                    selected_cats = [c.strip() for c in scope.categories.split(",") if c.strip()]
                    for row in t3.rows:
                        for cell in row.cells:
                            txt = cell.text.strip()
                            clean_txt = txt.replace('☒', '').replace('☐', '').replace('[X]', '').replace('[ ]', '').strip()
                            if clean_txt:
                                if clean_txt in selected_cats:
                                    cell.text = f"☒ {clean_txt}"
                                    for p in cell.paragraphs:
                                        for r in p.runs:
                                            r.bold = True
                                else:
                                    cell.text = f"☐ {clean_txt}"
                    body.insert(insert_idx, t3._tbl)
                    insert_idx += 1

                    p_sp1 = doc.add_paragraph()
                    body.insert(insert_idx, p_sp1._p)
                    insert_idx += 1

                    # Permits Checklist (Table 4)
                    t4 = docx.table.Table(copy.deepcopy(tbl4_xml), doc)
                    selected_permits = [p.strip() for p in scope.permits.split(",") if p.strip()]
                    for row in t4.rows:
                        for cell in row.cells:
                            txt = cell.text.strip()
                            clean_txt = txt.replace('☒', '').replace('☐', '').replace('[X]', '').replace('[ ]', '').strip()
                            if clean_txt:
                                if clean_txt in selected_permits:
                                    cell.text = f"☒ {clean_txt}"
                                    for p in cell.paragraphs:
                                        for r in p.runs:
                                            r.bold = True
                                else:
                                    cell.text = f"☐ {clean_txt}"
                    body.insert(insert_idx, t4._tbl)
                    insert_idx += 1

                    # Drawings Link
                    if p_draw_xml:
                        p_draw = docx.text.paragraph.Paragraph(copy.deepcopy(p_draw_xml), doc)
                        body.insert(insert_idx, p_draw._p)
                        insert_idx += 1

                    if p_folder_xml:
                        p_folder = docx.text.paragraph.Paragraph(copy.deepcopy(p_folder_xml), doc)
                        for run in p_folder.runs:
                            if "See folder 4." in run.text:
                                run.text = f"See folder 4.{idx+1}"
                        body.insert(insert_idx, p_folder._p)
                        insert_idx += 1

                    # Job Description table (Table 5)
                    t5 = docx.table.Table(copy.deepcopy(tbl5_xml), doc)
                    desc_cell = t5.cell(1, 0)
                    desc_cell.text = scope.description or "No description provided."
                    
                    if scope.image_path:
                        scope_img_bytes = get_image_bytes(scope.image_path)
                        if scope_img_bytes:
                            p_img = desc_cell.add_paragraph()
                            p_img.text = "\n"
                            try:
                                p_img.add_run().add_picture(io.BytesIO(scope_img_bytes), width=Inches(3.5))
                            except Exception as img_err:
                                logger.error(f"Error placing image in template cell: {str(img_err)}")

                    ref_cell = t5.cell(1, 1)
                    ref_cell.text = f"4.{idx+1}"

                    body.insert(insert_idx, t5._tbl)
                    insert_idx += 1

                    p_sp2 = doc.add_paragraph()
                    body.insert(insert_idx, p_sp2._p)
                    insert_idx += 1

                    # Subcontractor Table (Table 6)
                    t6 = docx.table.Table(copy.deepcopy(tbl6_xml), doc)
                    body.insert(insert_idx, t6._tbl)
                    insert_idx += 1

                    # Page Break
                    if idx < len(scopes_list) - 1:
                        p_pb = doc.add_paragraph()
                        p_pb.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
                        body.insert(insert_idx, p_pb._p)
                        insert_idx += 1

        # Delete empty paragraphs on the cover page to compensate for logo and cover image heights
        # This keeps the vesselName title at the top right and pulls the revision table back to page 1
        to_delete = []
        for idx, p in enumerate(doc.paragraphs[:15]):
            if p.text.strip() == "" and len(p.runs) == 0:
                has_drawing = "w:drawing" in p._element.xml or "w:pict" in p._element.xml
                if not has_drawing:
                    to_delete.append(p)
        for p in to_delete:
            p_element = p._element
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)

        docx_filename = f"Generated_Tender_Specification_Document_{report.report_id}.docx"
        
        if use_azure:
            try:
                logger.info(f"Saving generated DOCX '{docx_filename}' to Azure Storage container: {settings.AZURE_STORAGE_CONTAINER_REPORTS}")
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                azure_storage.upload_blob(
                    container_name=settings.AZURE_STORAGE_CONTAINER_REPORTS,
                    blob_name=docx_filename,
                    data=docx_buffer.getvalue()
                )
            except Exception as e:
                logger.error(f"Failed to upload generated DOCX to Azure: {e}. Saving locally as fallback.")
                docx_path = os.path.join(UPLOAD_DIR, docx_filename)
                doc.save(docx_path)
        else:
            docx_path = os.path.join(UPLOAD_DIR, docx_filename)
            doc.save(docx_path)

        return {
            "status": "success",
            "report_id": str(report.report_id),
            "message": "Tender Specification Document generated successfully!"
        }
    except Exception as e:
        db.rollback()
        logger.error(f"Error generating survey report: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")

@router.get("/generated/{report_id}/download-docx")
def download_docx(report_id: str):
    """
    Downloads the generated survey report DOCX file.
    """
    docx_filename = f"Generated_Tender_Specification_Document_{report_id}.docx"
    use_azure = bool(settings.AZURE_STORAGE_CONNECTION_STRING or (settings.AZURE_STORAGE_ACCOUNT_NAME and settings.AZURE_STORAGE_ACCOUNT_KEY))
    
    if use_azure:
        try:
            logger.info(f"Downloading generated DOCX '{docx_filename}' from Azure Storage container: {settings.AZURE_STORAGE_CONTAINER_REPORTS}")
            docx_bytes = azure_storage.download_blob(settings.AZURE_STORAGE_CONTAINER_REPORTS, docx_filename)
            return StreamingResponse(
                io.BytesIO(docx_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={docx_filename}"}
            )
        except Exception as e:
            logger.error(f"Failed to download generated DOCX from Azure: {e}. Trying local fallback.")

    # Local fallback
    file_path = os.path.join(UPLOAD_DIR, docx_filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Generated Word report not found")
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=docx_filename)

@router.get("/generated/{report_id}/download-pdf")
def download_pdf(report_id: str, db: Session = Depends(get_db)):
    """
    Generates and downloads a ReportLab PDF version mirroring the Word Template layout.
    """
    try:
        report_uuid = uuid.UUID(report_id)
        report = db.query(SurveyReport).filter(SurveyReport.report_id == report_uuid).first()
        if not report:
            raise HTTPException(status_code=404, detail="Survey report not found")

        project = db.query(Project).filter(Project.project_id == report.project_id).first()
        subcontractors = db.query(SurveyReportSubcontractor).filter(SurveyReportSubcontractor.report_id == report_uuid).all()
        scopes = db.query(SurveyReportWorkScope).filter(SurveyReportWorkScope.report_id == report_uuid).order_by(SurveyReportWorkScope.sequence_number).all()

        pdf_filename = f"Generated_Tender_Specification_Document_{report_id}.pdf"
        pdf_path = os.path.join(UPLOAD_DIR, pdf_filename)
        
        use_azure = bool(settings.AZURE_STORAGE_CONNECTION_STRING or (settings.AZURE_STORAGE_ACCOUNT_NAME and settings.AZURE_STORAGE_ACCOUNT_KEY))
        pdf_buffer = io.BytesIO() if use_azure else None

        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.pdfgen import canvas

        # Page numbering canvas helper
        class NumberedCanvas(canvas.Canvas):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self._saved_page_states = []

            def showPage(self):
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                num_pages = len(self._saved_page_states)
                for state in self._saved_page_states:
                    self.__dict__.update(state)
                    self.draw_page_elements(num_pages)
                    super().showPage()
                super().save()

            def draw_page_elements(self, page_count):
                self.saveState()
                if self._pageNumber > 1:
                    # Draw Running Header
                    self.setStrokeColor(colors.HexColor("#000000"))
                    self.setLineWidth(0.5)
                    self.line(40, 770, 572, 770)

                    self.setFont("Helvetica-Bold", 8)
                    self.setFillColor(colors.HexColor("#333333"))
                    self.drawString(40, 755, "Document:")
                    self.drawString(380, 755, "Page:")
                    self.drawString(380, 742, "Revision:")
                    self.drawString(40, 729, "Doc. Nr:")
                    self.drawString(380, 729, "Reference:")

                    self.setFont("Helvetica-Bold", 8)
                    self.drawString(100, 755, report.document_name or "")
                    self.drawString(100, 729, report.doc_nr or "")

                    self.setFont("Helvetica", 8)
                    self.drawString(440, 755, f"{self._pageNumber} of {page_count}")
                    self.drawString(440, 742, report.revision or "")
                    self.drawString(440, 729, report.reference or "")

                    self.line(40, 720, 572, 720)

                    # Draw Running Footer
                    self.setStrokeColor(colors.HexColor("#000000"))
                    self.setLineWidth(0.5)
                    self.line(40, 45, 572, 45)

                    self.setFont("Helvetica", 8)
                    self.setFillColor(colors.HexColor("#333333"))
                    self.drawString(40, 30, "Survey Report")
                    
                    current_date_str = datetime.date.today().strftime('%d %B %Y')
                    self.drawRightString(572, 30, f"Date: {current_date_str}")
                self.restoreState()

        doc = SimpleDocTemplate(pdf_buffer if use_azure else pdf_path, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=85, bottomMargin=45)
        story = []
        styles = getSampleStyleSheet()

        # Custom Styles
        title_style = ParagraphStyle(
            'CoverTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=28, leading=34, textColor=colors.HexColor('#1e3a8a'), spaceAfter=15
        )
        subtitle_style = ParagraphStyle(
            'CoverSubtitle', parent=styles['Normal'], fontName='Helvetica', fontSize=16, leading=20, textColor=colors.HexColor('#475569'), spaceAfter=80
        )
        h1_style = ParagraphStyle(
            'SectionH1', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=18, leading=22, textColor=colors.HexColor('#1e3a8a'), spaceAfter=12, keepWithNext=True
        )
        body_style = ParagraphStyle(
            'ReportBody', parent=styles['BodyText'], fontName='Helvetica', fontSize=10, leading=14, textColor=colors.HexColor('#334155'), spaceAfter=10
        )
        body_bold = ParagraphStyle(
            'ReportBodyBold', parent=body_style, fontName='Helvetica-Bold'
        )

        # ---------------- PAGE 1: COVER PAGE ----------------
        story.append(Spacer(1, 40))
        story.append(Paragraph(f"{report.company_name}", ParagraphStyle('Company', fontName='Helvetica-Bold', fontSize=16, textColor=colors.HexColor('#1e3a8a'))))
        story.append(Spacer(1, 60))
        story.append(Paragraph(f"{report.vessel_name}", title_style))
        story.append(Paragraph("Intermediate Survey Report 2026", subtitle_style))

        if report.cover_image_path:
            cover_bytes = get_image_bytes(report.cover_image_path)
            if cover_bytes:
                try:
                    story.append(Image(io.BytesIO(cover_bytes), width=280, height=210))
                    story.append(Spacer(1, 20))
                except Exception as pdf_img_err:
                    logger.error(f"Error placing cover image in PDF: {str(pdf_img_err)}")

        meta_data = [
            [Paragraph("<b>Document Description:</b>", body_style), Paragraph("Maintenance & Repair Survey Draft Report", body_style)],
            [Paragraph("<b>Created By:</b>", body_style), Paragraph("Chief Engineer", body_style)],
            [Paragraph("<b>Status:</b>", body_style), Paragraph(f"{report.status}", body_style)],
            [Paragraph("<b>Exported Date:</b>", body_style), Paragraph(f"{datetime.date.today().strftime('%d %B %Y')}", body_style)]
        ]
        meta_table = Table(meta_data, colWidths=[150, 300])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
            ('PADDING', (0,0), (-1,-1), 8),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1'))
        ]))
        story.append(meta_table)
        story.append(PageBreak())

        # ---------------- PAGE 2: GENERAL SCHEDULE & SUBCONTRACTORS ----------------
        story.append(Paragraph("1) General & Revision Details", h1_style))
        general_text = (
            "Disclaimer: Document may be subject to change. Check the latest version on eDOCS or intranet. "
            "This intermediate survey report compiles inspection photos and physical reports submitted by the "
            "vessel command team (Captain and Chief Engineer) during dry-docking preparation. AI Vision models "
            "have been utilized to generate initial shipyard work scopes which have been verified by the engineering team."
        )
        story.append(Paragraph(general_text, body_style))
        story.append(Spacer(1, 15))

        if report.interior_image_path:
            interior_bytes = get_image_bytes(report.interior_image_path)
            if interior_bytes:
                story.append(Paragraph("<b>Vessel Interior Capacity Layout:</b>", body_bold))
                try:
                    story.append(Image(io.BytesIO(interior_bytes), width=280, height=160))
                    story.append(Spacer(1, 15))
                except Exception as pdf_img_err:
                    logger.error(f"Error placing interior image in PDF: {str(pdf_img_err)}")

        story.append(Paragraph("2) Drydock Schedule Details", h1_style))
        sched_data = [
            [Paragraph("<b>Vessel Name:</b>", body_style), Paragraph(f"{report.vessel_name}", body_style)],
            [Paragraph("<b>ETA at Shipyard:</b>", body_style), Paragraph(f"{report.arrival_date.strftime('%Y-%m-%d') if report.arrival_date else ''}", body_style)],
            [Paragraph("<b>Expected Duration:</b>", body_style), Paragraph(f"{report.total_lead_time} Days", body_style)],
            [Paragraph("<b>Drydock Duration:</b>", body_style), Paragraph(f"{report.drydock_duration} Days", body_style)]
        ]
        sched_table = Table(sched_data, colWidths=[150, 300])
        sched_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f1f5f9')),
            ('PADDING', (0,0), (-1,-1), 8),
            ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0'))
        ]))
        story.append(sched_table)
        story.append(PageBreak())

        # ---------------- PAGE 3: INDEX OF SCOPES ----------------
        story.append(Paragraph("3) Index of Work Items", h1_style))
        story.append(Spacer(1, 10))

        index_data = [[
            Paragraph("<b>Item Nr.</b>", body_style), 
            Paragraph("<b>Shipyard Scope Title</b>", body_style), 
            Paragraph("<b>Tender Ref</b>", body_style), 
            Paragraph("<b>Status</b>", body_style),
            Paragraph("<b>Page No.</b>", body_style)
        ]]
        for idx, scope in enumerate(scopes):
            index_data.append([
                Paragraph(f"{idx+1}", body_style),
                Paragraph(scope.title, body_style),
                Paragraph(f"4.{idx+1}", body_style),
                Paragraph(scope.status or "New", body_style),
                Paragraph(str(4 + idx), body_style) # detailed scope idx starts on page 4
            ])

        index_table = Table(index_data, colWidths=[50, 230, 80, 80, 60])
        index_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e3a8a')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('PADDING', (0,0), (-1,-1), 6)
        ]))
        for i in range(5):
            index_table.setStyle(TableStyle([('TEXTCOLOR', (i,0), (i,0), colors.white)]))
        story.append(index_table)
        story.append(PageBreak())

        # ---------------- PAGE 4+: DETAILED WORK SCOPES ----------------
        story.append(Paragraph("4) Scope of Works Yard Details", h1_style))
        story.append(Spacer(1, 10))

        for idx, scope in enumerate(scopes):
            story.append(Paragraph(f"4.{idx+1} {scope.title}", ParagraphStyle('H2', fontName='Helvetica-Bold', fontSize=12, textColor=colors.HexColor('#0369a1'))))
            story.append(Spacer(1, 5))

            # 1. Categories checklist Table
            categories_list = ["Steel", "Piping", "Cleaning", "Transport", "Mechanical", "Electrical", "Painting", "Hydraulic"]
            checked_cats = [c.strip() for c in scope.categories.split(",") if c.strip()]
            cat_table_data = [
                [
                    Paragraph(f"{'<b>[X]</b>' if 'Steel' in checked_cats else '[ ]'} Steel", body_style),
                    Paragraph(f"{'<b>[X]</b>' if 'Piping' in checked_cats else '[ ]'} Piping", body_style),
                    Paragraph(f"{'<b>[X]</b>' if 'Cleaning' in checked_cats else '[ ]'} Cleaning", body_style),
                    Paragraph(f"{'<b>[X]</b>' if 'Transport' in checked_cats else '[ ]'} Transport", body_style)
                ],
                [
                    Paragraph(f"{'<b>[X]</b>' if 'Mechanical' in checked_cats else '[ ]'} Mechanical", body_style),
                    Paragraph(f"{'<b>[X]</b>' if 'Electrical' in checked_cats else '[ ]'} Electrical", body_style),
                    Paragraph(f"{'<b>[X]</b>' if 'Painting' in checked_cats else '[ ]'} Painting", body_style),
                    Paragraph(f"{'<b>[X]</b>' if 'Hydraulic' in checked_cats else '[ ]'} Hydraulic", body_style)
                ]
            ]
            cat_table = Table(cat_table_data, colWidths=[125, 125, 125, 125])
            cat_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
                ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
                ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#f1f5f9')),
                ('PADDING', (0,0), (-1,-1), 4)
            ]))
            story.append(cat_table)
            story.append(Spacer(1, 10))

            # 2. Detailed Job Table
            job_desc_p = Paragraph(scope.description or "No description provided.", body_style)
            img_flowable = None
            if scope.image_path:
                scope_img_bytes = get_image_bytes(scope.image_path)
                if scope_img_bytes:
                    try:
                        img_flowable = Image(io.BytesIO(scope_img_bytes), width=220, height=140)
                    except Exception as img_err:
                        logger.error(f"Error rendering scope image in PDF: {str(img_err)}")

            desc_cell_flowables = [job_desc_p]
            if img_flowable:
                desc_cell_flowables.append(Spacer(1, 10))
                desc_cell_flowables.append(img_flowable)

            desc_ref_p = Paragraph(f"Tender Ref: <b>4.{idx+1}</b>", body_style)

            detail_data = [
                [Paragraph("<b>Job Description for Yard:</b>", body_bold), Paragraph("<b>Tender reference:</b>", body_bold)],
                [desc_cell_flowables, desc_ref_p]
            ]
            detail_table = Table(detail_data, colWidths=[370, 130])
            detail_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#e2e8f0')),
                ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#f8fafc')),
                ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#94a3b8')),
                ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
                ('PADDING', (0,0), (-1,-1), 8),
                ('VALIGN', (0,0), (-1,-1), 'TOP')
            ]))
            story.append(detail_table)
            story.append(Spacer(1, 15))

            if idx < len(scopes) - 1:
                story.append(PageBreak())

        # Build Document
        if use_azure:
            try:
                doc.build(story, canvasmaker=NumberedCanvas)
                pdf_data = pdf_buffer.getvalue()
                logger.info(f"Uploading generated PDF '{pdf_filename}' to Azure Storage container: {settings.AZURE_STORAGE_CONTAINER_REPORTS}")
                azure_storage.upload_blob(
                    container_name=settings.AZURE_STORAGE_CONTAINER_REPORTS,
                    blob_name=pdf_filename,
                    data=pdf_data
                )
                return StreamingResponse(
                    io.BytesIO(pdf_data),
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"attachment; filename={pdf_filename}"}
                )
            except Exception as e:
                logger.error(f"Failed to compile and upload PDF to Azure: {e}. Building locally as fallback.")
                local_doc = SimpleDocTemplate(pdf_path, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=85, bottomMargin=45)
                local_doc.build(story, canvasmaker=NumberedCanvas)
                return FileResponse(pdf_path, media_type="application/pdf", filename=pdf_filename)
        else:
            doc.build(story, canvasmaker=NumberedCanvas)
            return FileResponse(pdf_path, media_type="application/pdf", filename=pdf_filename)
    except Exception as e:
        logger.error(f"Error compiling survey report PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"PDF compilation failed: {str(e)}")
